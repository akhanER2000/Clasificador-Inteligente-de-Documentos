"""
Generador de archivos Word (.docx) a partir de datos estructurados.
Usa python-docx para crear documentos formateados.
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_word(extracted_data: dict, document_type: str, output_dir: str = ".tmp") -> str:
    """
    Genera un archivo Word a partir de datos extraídos.
    
    Args:
        extracted_data: Dict con sections y metadata del documento.
        document_type: Tipo de documento clasificado.
        output_dir: Directorio de salida.
        
    Returns:
        Ruta al archivo generado.
    """
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{document_type}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    title = extracted_data.get("title", "Documento")
    metadata = extracted_data.get("metadata", {})
    content_sections = extracted_data.get("sections", [])
    
    # Título principal
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata como tabla de información
    if metadata:
        doc.add_paragraph()  # Espacio
        meta_table = doc.add_table(rows=0, cols=2)
        meta_table.style = "Light Shading Accent 1"
        
        for key, value in metadata.items():
            if value:
                row = meta_table.add_row()
                row.cells[0].text = key.capitalize()
                row.cells[1].text = str(value)
                # Negrita para la clave
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        doc.add_paragraph()  # Espacio después de metadata
    
    # Línea separadora
    doc.add_paragraph("_" * 60)
    
    # Contenido por secciones
    if content_sections:
        for section_data in content_sections:
            heading = section_data.get("heading")
            content = section_data.get("content", "")
            level = section_data.get("level", 1)
            
            # Agregar heading si existe
            if heading:
                heading_level = min(level, 4)  # Máximo nivel 4 en Word
                doc.add_heading(heading, level=heading_level)
            
            # Agregar contenido
            if content:
                # Detectar si hay listas (líneas que empiezan con - o •)
                lines = content.split("\n")
                current_paragraph = ""
                
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith(("-", "•", "*")):
                        # Si hay párrafo acumulado, escribirlo primero
                        if current_paragraph:
                            doc.add_paragraph(current_paragraph)
                            current_paragraph = ""
                        # Agregar como item de lista
                        bullet_text = stripped.lstrip("-•* ").strip()
                        doc.add_paragraph(bullet_text, style="List Bullet")
                    elif stripped.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                        if current_paragraph:
                            doc.add_paragraph(current_paragraph)
                            current_paragraph = ""
                        numbered_text = stripped.split(".", 1)[1].strip() if "." in stripped else stripped
                        doc.add_paragraph(numbered_text, style="List Number")
                    elif stripped == "":
                        if current_paragraph:
                            doc.add_paragraph(current_paragraph)
                            current_paragraph = ""
                    else:
                        if current_paragraph:
                            current_paragraph += " " + stripped
                        else:
                            current_paragraph = stripped
                
                # Escribir último párrafo acumulado
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
    else:
        doc.add_paragraph("No se pudo extraer contenido estructurado del documento.")
    
    # Pie de página con timestamp
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"Generado automáticamente el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(filepath)
    return filepath
